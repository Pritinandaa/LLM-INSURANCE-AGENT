import streamlit as st
import os
import pandas as pd
import numpy as np
from dotenv import load_dotenv
import plotly.express as px
import plotly.graph_objects as go
 
# LangChain & AI Imports
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from io import BytesIO
from docx import Document
 
# --- INITIALIZE ENVIRONMENT ---
load_dotenv()
GOOGLE_API_KEY = "AIzaSyCXurwAMAHUQcl14-E3q-i-5tpwNgzuhv8"
 
# --- PAGE CONFIG ---
st.set_page_config(page_title="Acuity AI - Financial Portfolio Advisor", layout="wide", initial_sidebar_state="expanded")

# --- CUSTOM CSS (Enhanced Corporate Dark Theme) ---
st.markdown("""
<style>
    .stApp { background-color: #0D1117; color: #E6E6E6; }
    [data-testid="stSidebar"] { background-color: #161B22; border-right: 1px solid #30363D; }
    
    /* Upload Box Styling */
    .upload-container {
        border: 2px dashed #30363D;
        padding: 40px;
        border-radius: 15px;
        text-align: center;
        background: rgba(255, 255, 255, 0.02);
        margin-top: 50px;
    }
    
    /* KPI Header */
    .kpi-container {
        display: flex;
        justify-content: space-around;
        padding: 15px;
        background: rgba(0, 150, 199, 0.1);
        border-radius: 12px;
        border: 1px solid #0096C7;
        margin-bottom: 25px;
    }
    .kpi-box { text-align: center; }
    .kpi-value { font-size: 1.8rem; font-weight: bold; color: #0096C7; }
    
    /* Info Boxes */
    .info-box {
        background: rgba(255, 255, 255, 0.05);
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #0096C7;
        margin-bottom: 15px;
    }
</style>
""", unsafe_allow_html=True)
 
# --- REPORT EXPORT ---
def build_report(messages):
    doc = Document()
    doc.add_heading("Acuity AI — Portfolio Intelligence Report", level=0)
    doc.add_paragraph(f"Generated from {len(messages)} conversation turns.")
    doc.add_paragraph("")
    for m in messages:
        role = "User" if m["role"] == "user" else "Acuity AI"
        doc.add_heading(role, level=2)
        doc.add_paragraph(m["content"])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state.messages = []
 
# --- SIDEBAR (Control Panel Only) ---
with st.sidebar:
    st.markdown("<h1 style='color: #0096C7;'>🛡️ Acuity AI</h1>", unsafe_allow_html=True)
    st.caption("Strategic Control Panel")
    st.divider()
    
    st.markdown("### 🖥️ Display Controls")
    show_heatmap = st.toggle("Risk Heatmap", value=False, help="Show a visual heatmap of risk scores across portfolio chunks.")
    show_gauge = st.toggle("Health Gauge", value=False, help="Display a gauge showing overall portfolio health.")
    show_table = st.toggle("Detailed Logs", value=False, help="View detailed risk analysis table.")
    
    st.divider()
    st.markdown("### ℹ️ About")
    st.info("Acuity AI analyzes your financial portfolio PDF to identify risks, provide insights, and offer AI-powered advice.")
    
    st.markdown("### 📋 How to Use")
    with st.expander("Quick Guide"):
        st.markdown("""
        1. **Upload PDF**: Upload your portfolio report (PDF format).
        2. **View Insights**: Check KPIs and visualizations.
        3. **Chat with AI**: Ask questions about risks, dependencies, or strategies.
        4. **Download Report**: Export your conversation as a Word document.
        """)
    
    if st.button("Reset Session", use_container_width=True, help="Clear all chat history and start fresh."):
        st.session_state.messages = []
        st.rerun()
 
# --- MAIN SCREEN LOGIC ---
 
# 1. Check for API Key first
if not GOOGLE_API_KEY:
    st.error("Missing GOOGLE_API_KEY. Please ensure your .env file is configured correctly.")
    st.stop()
 
# 2. Main Title and Intro
st.markdown("<h1 style='text-align: center; color: #0096C7;'>🛡️ Acuity AI - Financial Portfolio Advisor</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #8B949E;'>Intelligent analysis of your investment portfolio using AI-powered insights and risk assessment.</p>", unsafe_allow_html=True)
st.divider()
 
# 3. File Upload (Conditional Rendering)
st.markdown("### 📄 Upload Your Portfolio Report")
uploaded_pdf = st.file_uploader("Upload Portfolio PDF", type="pdf", label_visibility="collapsed", help="Select a PDF file containing your financial portfolio details for analysis.")
 
if not uploaded_pdf:
    # Display Welcome Screen if no file is present
    st.markdown("""
    <div class="upload-container">
        <h2 style='color: #E6E6E6;'>Welcome to Acuity AI</h2>
        <p style='color: #8B949E;'>To begin your intelligence analysis, please upload your project portfolio or executive PDF report above.</p>
        <div class="info-box">
            <strong>What to expect:</strong> Our AI will analyze your portfolio for risks, provide health scores, and answer your questions about investments, dependencies, and strategies.
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    st.success("PDF uploaded successfully! Analyzing your portfolio...")
    # --- PROCESSING LOGIC (Runs once per file) ---
    @st.cache_resource
    def process_data(file):
        with open("active_portfolio.pdf", "wb") as f:
            f.write(file.getbuffer())
        
        loader = PyPDFLoader("active_portfolio.pdf")
        pages = loader.load()
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = splitter.split_documents(pages)
        
        # Simple Risk Scoring logic
        RISK_MAP = {"delay": 2, "over budget": 3, "critical path": 2, "risk": 1, "slippage": 2}
        risk_list = []
        for i, d in enumerate(docs):
            s = min(sum(v for k, v in RISK_MAP.items() if k in d.page_content.lower()), 10)
            risk_list.append({"Chunk": i, "Score": s, "Content": d.page_content[:150]})
        
        df = pd.DataFrame(risk_list)
        
        # LLM & Vector Store
        llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
        vector = FAISS.from_documents(docs, embeddings)
        
        return df, vector.as_retriever(), llm
 
    with st.spinner("Analyzing Portfolio Architecture..."):
        risk_df, retriever, llm = process_data(uploaded_pdf)
    
    # Calculate Metrics
    avg_score = risk_df["Score"].mean()
    health_idx = max(0, 100 - (avg_score * 12))
 
    # --- ACTIVE DASHBOARD UI ---
    
    st.markdown("### 📊 Portfolio Analysis Dashboard")
    st.markdown("Below is a summary of your portfolio's health based on AI analysis of the uploaded document.")
    
    # KPI Bar
    st.markdown(f"""
    <div class="kpi-container">
        <div class="kpi-box">Portfolio Health<br><span class="kpi-value">{health_idx:.0f}%</span><br><small>Higher is better</small></div>
        <div class="kpi-box">Avg Risk Score<br><span class="kpi-value">{avg_score:.1f}</span><br><small>Out of 10</small></div>
        <div class="kpi-box">Analysis Chunks<br><span class="kpi-value">{len(risk_df)}</span><br><small>Document sections</small></div>
    </div>
    """, unsafe_allow_html=True)
 
    # Layout: Split vs Single
    any_viz = any([show_heatmap, show_gauge, show_table])
    
    if any_viz:
        col_viz, col_chat = st.columns([1, 1], gap="large")
    else:
        col_chat = st.container()
 
    # VISUALS COLUMN
    if any_viz:
        with col_viz:
            st.subheader("📊 Analytical Overlays")
            st.markdown("Visual representations of your portfolio's risk profile.")
            if show_heatmap:
                st.markdown("**Risk Heatmap:** Shows risk scores across different sections of your portfolio document.")
                fig = px.imshow(np.array([risk_df["Score"]]), color_continuous_scale="Reds")
                fig.update_layout(height=200, margin=dict(l=0, r=0, t=10, b=0))
                st.plotly_chart(fig, use_container_width=True)
            
            if show_gauge:
                st.markdown("**Health Gauge:** Overall portfolio health indicator.")
                gauge = go.Figure(go.Indicator(
                    mode="gauge+number", value=health_idx,
                    gauge={'axis': {'range': [0, 100]}, 'bar': {'color': "#0096C7"}}))
                gauge.update_layout(height=250, margin=dict(l=10, r=10, t=40, b=10))
                st.plotly_chart(gauge, use_container_width=True)
            
            if show_table:
                st.markdown("**Detailed Risk Analysis:** Breakdown of risk scores for each document chunk.")
                st.dataframe(risk_df, hide_index=True, use_container_width=True)
 
    # CHAT COLUMN
    with (col_chat if any_viz else st.container()):
        st.subheader("💬 AI Financial Advisor")
        st.markdown("Ask questions about your portfolio, risks, investment strategies, or get personalized advice.")
        
        chat_container = st.container(height=550 if any_viz else 650)
        with chat_container:
            for m in st.session_state.messages:
                with st.chat_message(m["role"]):
                    st.markdown(m["content"])
 
        if prompt := st.chat_input("E.g., 'What are the main risks in my portfolio?' or 'Suggest diversification strategies.'"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with chat_container:
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                with st.chat_message("assistant"):
                    # RAG Context Retrieval
                    context_docs = retriever.invoke(prompt)
                    context_text = "\n".join([d.page_content for d in context_docs])
                    
                    full_p = f"System: You are Acuity AI, a Portfolio Analyst. Use context.\nContext: {context_text}\nUser: {prompt}"
                    response = llm.invoke(full_p)
                    st.markdown(response.content)
            
            st.session_state.messages.append({"role": "assistant", "content": response.content})

        # Download Report Button
        if st.session_state.messages:
            st.markdown("### 📥 Export Conversation")
            st.download_button(
                label="Download Report as Word Document",
                data=build_report(st.session_state.messages),
                file_name="AcuityAI_Portfolio_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Export your chat conversation and analysis as a downloadable report."
            )

st.divider()
st.markdown("<p style='text-align: center; color: #8B949E;'>Powered by Acuity AI | Multi-Agent Financial Analysis System</p>", unsafe_allow_html=True)
 
